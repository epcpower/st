import collections
import itertools
import logging
import os
import signal

import attr
import canmatrix.formats
import click
import docx
import docx.enum.section
import docx.enum.table
import docx.enum.text
import docx.shared
import docx.table
import lxml.etree

import epyqlib.utils.general

# See file COPYING in this source tree
__copyright__ = 'Copyright 2017, EPC Power Corp.'
__license__ = 'GPLv2+'


logger = logging.getLogger(__name__)


def docx_ancestor(element, target_type=docx.document.Document):
    while not isinstance(element, target_type) and element is not None:
        element = element._parent

    return element


def w(s):
    return '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}' + s


def shd_element(fill):
    e = lxml.etree.Element(w('shd'))
    # <w:shd w:val="clear" w:color="auto" w:fill="D9D9D9" w:themeFill="background1" w:themeFillShade="D9"/>

    e.attrib[w('fill')] = fill

    return e


def shade(cell, fill):
    cell._tc.tcPr.append(shd_element(fill=fill))


@attr.s
class Table:
    title = attr.ib()
    headings = attr.ib()
    widths = attr.ib()
    total_width = attr.ib(default=10)
    comment = attr.ib(default='')
    rows = attr.ib(default=attr.Factory(list))

    def fill_docx(self, table):
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        # table.autofit = True
        # table.style = 'CAN Table Base'

        row = table.add_row()
        row.cells[0].text = self.title
        title_paragraph = row.cells[0].paragraphs[0]
        title_paragraph.paragraph_format.keep_with_next = True
        title_paragraph.style = 'CAN Table Title'
        shade(row.cells[0], fill="000000")

        if len(self.comment) > 0:
            row = table.add_row()
            row.cells[0].merge(row.cells[-1])
            row.cells[0].text = self.comment
            row.cells[0].paragraphs[0].paragraph_format.keep_with_next = True

        row = table.add_row()
        for cell, heading in zip(row.cells, self.headings):
            cell.text = heading
            cell.paragraphs[0].style = 'CAN Table Heading'
            cell.paragraphs[0].paragraph_format.keep_with_next = True

        shadings = (
            {'fill': 'D9D9D9'},
            None,
        )

        for r, shading in zip(self.rows, itertools.cycle(shadings)):
            row = table.add_row()
            for cell, text in zip(row.cells, r):
                cell.text = str(text)
                if shading is not None:
                    shade(cell, **shading)
                cell.paragraphs[0].style = 'CAN Table Contents'
                cell.paragraphs[0].paragraph_format.keep_with_next = True

        remaining_width = (
            self.total_width - sum(w for w in self.widths if w is not None)
        )
        each_width = remaining_width / sum(1 for w in self.widths if w is None)
        widths = [each_width if w is None else w for w in self.widths]
        widths = [
            w if w is None else docx.shared.Inches(w)
            for w in widths
        ]
        for column, width in zip(table.columns, widths):
            column.width = width
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = width

        table.rows[0].cells[0].merge(table.rows[0].cells[-1])


def id_string(id):
    return '0x{:08X}'.format(id)


def tabulate_signals(signals):
    rows = []

    for signal in signals:
        # if len(signal.unit) == 0 and signal.factor != 1:
            # table.append(('***',) * 5)
        startbit = signal.getStartbit()
        rows.append((
            '', '',
            signal.name,
            '{}/{}'.format(startbit % 8, startbit // 8),
            signal.signalsize,
            signal.factor,
            signal.unit,
            '{}: {}'.format(signal.enumeration, signal.values) if
            signal.enumeration is not None else '',
        ))
        if len(signal.comment) > 0:
            rows.append((*(('',) * 7), signal.comment))

    return rows


def doc_signals(signals):
    rows = []

    for signal in signals:
        startbit = signal.getStartbit()
        enumeration = (
            signal.enumeration
            if signal.enumeration is not None
            else ''
        )
        rows.append((
            signal.name,
            '{}/{}'.format(startbit % 8, startbit // 8),
            signal.signalsize,
            signal.factor,
            signal.unit,
            enumeration,
            signal.comment,
        ))

    return rows

@click.command()
@click.option('--can', '-c', type=click.File('rb'), required=True)
@click.option('--template', '-t', type=click.File('rb'), required=True)
@click.option('--output', '-o', type=click.File('wb'), required=True)
@click.option('--verbose', '-v', count=True)
def main(can, template, output, verbose):
    if verbose >= 1:
        logger.setLevel(logging.DEBUG)

    matrix, = canmatrix.formats.load(
        can,
        os.path.splitext(can.name)[1].lstrip('.')
    ).values()

    mux_table_header = (
        'Mux Name', 'Mux ID', 'Name', 'Start', 'Length', 'Factor', 'Units',
        'Enumeration', 'Comment'
    )
    mux_table = epyqlib.utils.general.TextTable()
    mux_table.append(mux_table_header)
    mux_table_header = mux_table_header[2:]

    frame_table_header = (
        'Frame', 'ID', 'Name', 'Start', 'Length', 'Factor', 'Units',
        'Enumeration', 'Comment'
    )

    frame_table = epyqlib.utils.general.TextTable()
    frame_table.append(frame_table_header)
    frame_table_header = frame_table_header[2:]

    widths = [0.625] * len(frame_table_header)
    widths[0] = 2
    widths[-2] = 1.5
    widths[-1] = None

    enum_table_header = ('Value', 'Name')

    frame_tables = []
    multiplex_tables = []
    enumeration_tables = []

    for frame in sorted(matrix.frames, key=lambda f: f.name):
        frame_table.append(
            frame.name,
            id_string(frame.id),
        )

        a_ft = Table(
            title='{} ({})'.format(frame.name, id_string(frame.id)),
            comment=frame.comment,
            headings=frame_table_header,
            widths=widths,
        )
        frame_tables.append(a_ft)

        mux_table.append(
            '{} ({})'.format(frame.name, id_string(frame.id)),
        )

        multiplex_signal = frame.signals[0]
        if multiplex_signal.multiplex is None:
            multiplex_signal = None

        if multiplex_signal is None:
            signals = sorted(frame.signals, key=lambda s: s.name)
            table = tabulate_signals(signals)
            frame_table.extend(table)

            a_ft.rows.extend(doc_signals(signals))
        else:
            signals = (multiplex_signal,)
            table = tabulate_signals(signals)
            frame_table.extend(table)
            a_ft.rows.extend(doc_signals(signals))
            # multiplex_signal.values = {
            #     int(k): v for k, v in multiplex_signal.values.items()
            # }

            for value, name in sorted(multiplex_signal.values.items()):
                if value == 0:
                    continue

                a_mt = Table(
                    title='{} ({}) - {} ({})'.format(
                        frame.name,
                        id_string(frame.id),
                        name,
                        value
                    ),
                    # comment=frame.comment,
                    headings=mux_table_header,
                    widths=widths,
                )
                multiplex_tables.append(a_mt)

                mux_table.append(
                    '',
                    value,
                    name,
                )

                signals = tuple(
                    s for s in frame.signals
                    if s.multiplex == value and s.name != 'ReadParam_command'
                )

                mux_table.extend(tabulate_signals(sorted(
                    signals, key=lambda s: s.name)))

                a_mt.rows.extend(doc_signals(signals))

    print('\n\n - - - - - - - Multiplexes\n')
    print(mux_table)

    print('\n\n - - - - - - - Frames\n')
    print(frame_table)

    enumeration_table = epyqlib.utils.general.TextTable()
    enumeration_table.append('Name', '', 'Value')
    widths = (1.5, None)
    for name, values in sorted(matrix.valueTables.items()):
        a_et = Table(
            title=name,
            headings=enum_table_header,
            widths=widths,
            total_width=5,
        )
        enumeration_tables.append(a_et)
        a_et.rows.extend(sorted(values.items()))

        enumeration_table.append(name)
        for i, s in sorted(values.items()):
            enumeration_table.append('', s, i)

    print('\n\n - - - - - - - Enumerations\n')
    print(enumeration_table)

    doc = docx.Document(template)

    table_sets = {
        'frames': frame_tables,
        'multiplexers': multiplex_tables,
        'enumerations': enumeration_tables,
    }
    for tag, tables in table_sets.items():
        full_tag = '<gen_{}>'.format(tag)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == full_tag:
                break
            elif full_tag in paragraph.text:
                f = 'Tag {} found, expected as only text in paragraph: {}'
                raise Exception(f.format(full_tag, repr(paragraph.text)))
        else:
            raise Exception('Tag not found: {}'.format(full_tag))

        for table in tables:
            doc_table = doc.add_table(rows=0, cols=len(table.headings))
            doc_paragraph = doc.add_paragraph()

            paragraph._p.addprevious(doc_table._tbl)
            paragraph._p.addprevious(doc_paragraph._p)

            table.fill_docx(doc_table)

        # TODO: Would rather delete the tag paragraph but that breaks the
        #       template's landscape page format for some reason
        # paragraph._p.getparent().remove(paragraph._p)
        paragraph.clear()
    doc.save(output)
    pass


def _entry_point():
    logging.basicConfig(format='%(asctime)s - %(levelname)s - %(message)s')

    signal.signal(signal.SIGINT, signal.SIG_DFL)

    return main()
